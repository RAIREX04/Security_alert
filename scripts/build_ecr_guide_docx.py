from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PANDUAN_APLIKASI_ECR.md"
OUT_DIR = ROOT / "output" / "docs"
OUT = OUT_DIR / "Panduan_Aplikasi_ECR.docx"


TOKENS = {
    "font": "Calibri",
    "body_size": 11,
    "body_after": 6,
    "line_spacing": 1.25,
    "blue": "005BAC",
    "navy": "0F172A",
    "red": "C81E29",
    "muted": "475569",
    "light_blue": "E8EEF5",
    "soft": "F8FAFC",
    "border": "CBD5E1",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD5E1", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_paragraph(paragraph, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = RGBColor.from_string(TOKENS["navy"])
    normal.paragraph_format.space_after = Pt(TOKENS["body_after"])
    normal.paragraph_format.line_spacing = TOKENS["line_spacing"]

    for style_name, size, color, before, after in [
        ("Heading 1", 16, TOKENS["blue"], 18, 10),
        ("Heading 2", 13, TOKENS["blue"], 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = TOKENS["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Panduan Aplikasi ECR")
    set_font(run, size=9, color="64748B")


def add_cover(doc):
    p = doc.add_paragraph()
    style_paragraph(p, after=2)
    run = p.add_run("PANDUAN PENGGUNAAN")
    set_font(run, size=13, color=TOKENS["blue"], bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, after=6, line_spacing=1.05)
    run = p.add_run("APLIKASI ECR")
    set_font(run, size=34, color=TOKENS["navy"], bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, after=18)
    run = p.add_run("Emergency Call Response")
    set_font(run, size=16, color=TOKENS["red"], bold=True)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, TOKENS["border"])
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(para, before=18, after=4)
    run = para.add_run("[GAMBAR COVER - letakkan screenshot mockup aplikasi ECR atau dashboard utama]")
    set_font(run, size=14, color=TOKENS["red"], bold=True)
    para = cell.add_paragraph("Catatan: ambil tampilan yang memperlihatkan identitas aplikasi ECR dan dashboard operasional.")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(para, after=18)
    set_font(para.runs[0], size=10, color=TOKENS["muted"])

    p = doc.add_paragraph()
    style_paragraph(p, before=14, after=6)
    run = p.add_run("Dokumen ini berisi panduan penggunaan aplikasi, alur role, dan daftar titik screenshot yang harus diambil sendiri untuk finalisasi manual.")
    set_font(run, size=11, color=TOKENS["muted"])

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [Inches(1.75), Inches(4.75)])
    for row in meta.rows:
        for cell in row.cells:
            set_cell_border(cell, "DCE6F5", "6")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("Aplikasi", "ECR - Emergency Call Response"),
        ("Domain", "Emergency / Security Alert"),
        ("Role", "User, Staff, Admin, Superadmin, View Only"),
        ("Output", "Panduan penggunaan dengan placeholder screenshot"),
    ]
    for idx, (label, value) in enumerate(rows):
        left, right = meta.rows[idx].cells
        set_cell_shading(left, TOKENS["light_blue"])
        left.text = label
        right.text = value
        for cell in (left, right):
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, after=0)
                for run in paragraph.runs:
                    set_font(run, size=10, color=TOKENS["navy"], bold=cell is left)

    doc.add_page_break()


def add_placeholder_box(doc, placeholder, note, caption):
    table = doc.add_table(rows=3, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "FFF1F2")
    set_cell_border(cell, "FDA4AF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=18, after=18)
    run = p.add_run(placeholder)
    set_font(run, size=12, color=TOKENS["red"], bold=True)

    cell = table.cell(1, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "D7E3F0", "6")
    p = cell.paragraphs[0]
    style_paragraph(p, before=4, after=4)
    run = p.add_run("Catatan screenshot: ")
    set_font(run, size=10, color=TOKENS["navy"], bold=True)
    run = p.add_run(note)
    set_font(run, size=10, color=TOKENS["muted"])

    cell = table.cell(2, 0)
    set_cell_shading(cell, "FFFFFF")
    set_cell_border(cell, "D7E3F0", "6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=3, after=3)
    run = p.add_run(caption)
    set_font(run, size=9.5, color=TOKENS["muted"], bold=True)

    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=6)


def add_step_paragraph(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    style_paragraph(p, after=4)
    run = p.add_run(f"{number}. ")
    set_font(run, size=10.5, color=TOKENS["navy"], bold=True)
    run = p.add_run(text)
    set_font(run, size=10.5, color=TOKENS["navy"])


def add_toc(doc, section_titles):
    doc.add_heading("Daftar Isi", level=1)
    for title in section_titles:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=2)
        run = p.add_run(title)
        set_font(run, size=10.5, color=TOKENS["navy"])
    doc.add_page_break()


def parse_markdown(lines):
    blocks = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            blocks.append(("step", re.sub(r"^\d+\.\s+", "", line).strip()))
        elif line.startswith("[GAMBAR") or line.startswith("[Gambar") or line.startswith("[GAMBAR COVER"):
            placeholder = line.strip()
            note = ""
            caption = ""
            idx += 1
            while idx < len(lines):
                next_line = lines[idx].strip()
                if not next_line:
                    idx += 1
                    continue
                if next_line.startswith("Catatan screenshot:"):
                    idx += 1
                    note_lines = []
                    while idx < len(lines):
                        candidate = lines[idx].strip()
                        if not candidate:
                            idx += 1
                            continue
                        if candidate.startswith("Caption:"):
                            break
                        if candidate.startswith("#") or candidate.startswith("[GAMBAR"):
                            idx -= 1
                            break
                        note_lines.append(candidate)
                        idx += 1
                    note = " ".join(note_lines)
                    continue
                if next_line.startswith("Caption:"):
                    idx += 1
                    caption_lines = []
                    while idx < len(lines):
                        candidate = lines[idx].strip()
                        if not candidate:
                            idx += 1
                            continue
                        if candidate.startswith("#") or candidate.startswith("[GAMBAR"):
                            idx -= 1
                            break
                        caption_lines.append(candidate)
                        idx += 1
                    caption = " ".join(caption_lines)
                    break
                if next_line.startswith("#") or next_line.startswith("[GAMBAR"):
                    idx -= 1
                    break
                idx += 1
            blocks.append(("placeholder", placeholder, note, caption))
        elif line.endswith(":") and len(line) < 80:
            blocks.append(("label", line))
        else:
            parts = [line]
            idx += 1
            while idx < len(lines):
                nxt = lines[idx].strip()
                if not nxt:
                    break
                if nxt.startswith("#") or nxt.startswith("[GAMBAR") or re.match(r"^\d+\.\s+", nxt) or nxt.endswith(":"):
                    idx -= 1
                    break
                parts.append(nxt)
                idx += 1
            blocks.append(("p", " ".join(parts)))
        idx += 1
    return blocks


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    blocks = parse_markdown(lines)
    section_titles = [block[1] for block in blocks if block[0] == "h1"]

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc, section_titles)

    step_counter = 0
    for block in blocks:
        kind = block[0]
        if kind == "title":
            continue
        if kind == "h1":
            doc.add_heading(block[1], level=1)
        elif kind == "h2":
            doc.add_heading(block[1], level=2)
        elif kind == "label":
            p = doc.add_paragraph()
            style_paragraph(p, before=2, after=3)
            run = p.add_run(block[1])
            set_font(run, size=11, color=TOKENS["blue"], bold=True)
            if block[1].lower().startswith("langkah penggunaan"):
                step_counter = 0
        elif kind == "step":
            step_counter += 1
            add_step_paragraph(doc, step_counter, block[1])
        elif kind == "placeholder":
            _kind, placeholder, note, caption = block
            add_placeholder_box(doc, placeholder, note or "Ambil screenshot sesuai nama halaman.", caption or "Gambar. Screenshot aplikasi ECR")
        elif kind == "p":
            p = doc.add_paragraph()
            style_paragraph(p, after=6)
            run = p.add_run(block[1])
            set_font(run, size=11, color=TOKENS["navy"])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_docx())
