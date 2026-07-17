from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image


ROOT = Path(r"E:\security_alert")
SOURCE_DOCX = Path(r"E:\Panduan_Aplikasi_ECR_Profesional.docx")
OUT_DOCX = Path(r"E:\Panduan_Aplikasi_ECR_Profesional_SecurityAlert_SS.docx")


SCREENSHOTS = {
    "COVER": "launch.png",
    1: "launch.png",
    2: "secure-launch3.png",
    3: "report-form-final.png",
    4: "report-form-final.png",
    5: "status-final.png",
    6: "history-viewport.png",
    7: "staff-dashboard-scroll.png",
    8: "staff-report-detail.png",
    9: "staff-report-detail-2.png",
    10: "staff-form-top.png",
    11: "admin-dashboard-top2.png",
    12: "admin-history.png",
    13: "admin-report-detail.png",
    14: "admin-history.png",
    15: "admin-dashboard-top2.png",
    16: "admin-dept-detail.png",
    17: "admin-users.png",
    18: "admin-user-detail.png",
    19: "admin-users.png",
    20: "admin-history.png",
    21: "admin-report-detail.png",
    22: "admin-profile.png",
    23: "admin-dashboard-top2.png",
}


def find_placeholder_key(text: str):
    upper = text.upper()
    if "PLACEHOLDER GAMBAR COVER" in upper or "[GAMBAR COVER" in upper:
        return "COVER"
    match = re.search(r"(?:\[)?GAMBAR\s+(\d+)\b", upper)
    if match:
        return int(match.group(1))
    return None


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def preferred_width(image_path: Path):
    with Image.open(image_path) as image:
        width, height = image.size
    if width >= height:
        return Inches(5.95)
    if height / max(width, 1) > 1.75:
        return Inches(2.8)
    return Inches(3.4)


def fill_document():
    doc = Document(SOURCE_DOCX)
    filled = []
    missing = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().lower() == "placeholder screenshot":
            clear_paragraph(paragraph)
            paragraph.add_run("Screenshot aplikasi")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                key = find_placeholder_key(cell_text)
                if key is None:
                    continue

                filename = SCREENSHOTS.get(key)
                if not filename:
                    missing.append(str(key))
                    continue

                image_path = ROOT / filename
                if not image_path.exists():
                    missing.append(f"{key}: {image_path}")
                    continue

                paragraph = cell.paragraphs[0]
                clear_paragraph(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=preferred_width(image_path))

                for extra in cell.paragraphs[1:]:
                    clear_paragraph(extra)

                filled.append((key, str(image_path)))

    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(f"filled={len(filled)}")
    if missing:
        print("missing=" + "; ".join(missing))


if __name__ == "__main__":
    fill_document()
